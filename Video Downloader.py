import yt_dlp as youtube_dl
import urllib.request
from PIL import Image
from docx import Document



class vid_downloader:

    def __init__(self, language, path, urls, vid_or_aud):
        self.language = language
        self.path = path
        self.urls = urls
        self.vid_or_aud = vid_or_aud
        self.info = [[],[]]

    def input_info(self, video_url):

        with youtube_dl.YoutubeDL() as ydl:

            video_info = ydl.extract_info(video_url, download=False)

            # Obtaining title, description, and thumbnail url
            title = video_info.get('title', None)
            description = video_info.get("description", None)

            self.info[0].append(title)
            self.info[1].append(description)
       
    def get_vid_info(self):
        return self.info

    def download_video(self, video_url, video_count, language):    
        
        # Download the video using youtube-dl
        # Program Configurations
        #   downloads in mp4
        #   downloads to specific folder
        #   provides general video information (title, description, etc.)
        ydl_opts = {
            'noplaylist': True,
            'format':'mp4',
            'outtmpl': self.path + f'Psych2go {language} Video {video_count}.%(ext)s' 
        }

        with youtube_dl.YoutubeDL(ydl_opts) as ydl:
            video_info = ydl.extract_info(video_url, download=False)
            
            # Obtaining title, description, and thumbnail url
            title = video_info.get('title', None)
            description = video_info.get("description", None)
            thumbnail_url = video_info.get('thumbnail', None)

            # Downloading Video
            print("\nDownloading vid...")
            ydl.download([video_url])
            print("Download completed!!")

            # Downloading thumbnail
            self.download_thumbnail(thumbnail_url, self.path, video_count, language)

            # Formatting Thumbnail
            self.edit_thumbnail(video_count, self.path, language)

            # Writing title and description to a document
            self.save_info(title, description, video_count, self.path)

    def download_audio(self, video_url, video_count, language):
        
        # Download the video using youtube-dl
        # Program Configurations
        #   downloads in mp3
        #   downloads to specific folder
        #   provides general video information (title, description, etc.)
        ydl_opts = {
            'noplaylist': True,
            'verbose': True,
            'format': 'bestaudio/best',
            'outtmpl': self.path + f'Psych2go {language} Audio {video_count}.%(ext)s',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }]  
        }

        with youtube_dl.YoutubeDL(ydl_opts) as ydl:

            # Downloading Video
            print("\nDownloading audio...")
            ydl.download([video_url])
            print("Download completed!!")

    def download_thumbnail(self, thumbnail_url, path, num, language):

        full_path = path + f"{language} Video Thumbnail {num}" + ".jpg"

        print("\nDownloading Thumbnail...")
        urllib.request.urlretrieve(thumbnail_url, full_path)
        print("Download completed!!")

    def edit_thumbnail(self, num, path, language):

        print("\nEditing Thumbnail...")
        image = Image.open(path + f"{language} Video Thumbnail {num}.jpg")
        image = image.resize((1000, 1050))
        image.save(path + f"{language} Video Thumbnail {num}.jpg")
        print("Thumbnail Edited!!")

    def save_info(self, title, description, num, path):
        
        doc = Document(path + 'Psych2goInfo.docx')
        
        print("\nUploading Info...")
        doc.add_paragraph(f"Title {num}:\n\n " + title + "\n\n")
        doc.add_paragraph(f"Description {num}:\n\n " + description + "\n\n")
        doc.add_page_break()
        print("Upload Finished!!")
        doc.save(path + 'Psych2goInfo.docx')

    def video_or_audio(self, vid_or_aud):
        if vid_or_aud == 1:
            for url in self.urls:
                video_count = self.urls.index(url) + 1
                self.download_video(url, video_count, self.language)
                self.input_info(url)

        elif vid_or_aud == 2:
            for url in self.urls:
                video_count = self.urls.index(url) + 1
                self.download_video(url, video_count, self.language)
                self.download_audio(url, video_count, self.language)
                self.input_info(url)
